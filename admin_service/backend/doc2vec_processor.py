from langchain.docstore.document import Document
from langchain_mongodb import MongoDBAtlasVectorSearch
from langchain_openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from datetime import datetime
from google.cloud import storage
import fitz
import re
import os
import pandas as pd
import time
import json
from docx import Document as docxDocument
from io import BytesIO
from openpyxl import load_workbook
import pdfplumber

# os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = os.path.abspath('utility-encoder-420001-0db7ee074ec6.json')
# # Instantiates a client
# storage_client = storage.Client()

def make_columns_unique(columns):
    """
    使列名唯一的辅助函数，为重复的列名添加后缀。
    
    参数:
        columns (list): 列名列表
    
    返回:
        list: 唯一的列名列表
    """
    seen = {}
    result = []
    for col in columns:
        if col in seen:
            seen[col] += 1
            new_col = f"{col}_{seen[col]}"
        else:
            seen[col] = 0
            new_col = col
        result.append(new_col)
    return result

def split_pdf(file_data, output_folder, pages_per_split=100):
    """
    使用 PyMuPDF 将 PDF 按页切割为多个小文件。

    参数:
        file_data (bytes): 输入 PDF 文件的二进制数据
        output_folder (str): 输出文件夹路径。
        pages_per_split (int): 每个文件的页数。
    """
    # 确保输出文件夹存在
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # 打开 PDF 文件
    pdf_file_io = BytesIO(file_data)
    doc = fitz.open(stream=pdf_file_io, filetype="pdf")
    total_pages = doc.page_count
    
    output_files = []
    
    # 分块处理 PDF，每块处理 pages_per_split 页
    for start_page in range(0, total_pages, pages_per_split):
        pdf_writer = fitz.open()  # 创建一个新的空 PDF 文档
        
        # 获取每一块的结束页
        end_page = min(start_page + pages_per_split, total_pages)
        
        for page_num in range(start_page, end_page):
            pdf_writer.insert_pdf(doc, from_page=page_num, to_page=page_num)
        
        # 定义输出文件的路径
        output_pdf_path = os.path.join(output_folder, f"split_{start_page+1}_to_{end_page}.pdf")
        
        # 保存新的 PDF 文件，使用垃圾回收来优化资源使用
        pdf_writer.save(output_pdf_path, garbage=4)
        pdf_writer.close()
        
        print(f"已保存 {output_pdf_path}")
        output_files.append(output_pdf_path)

    # 关闭原始文档
    doc.close()
    
    return output_files

def clean_repeated_characters(text):
    # 使用正则表达式移除连续重复的字符
    cleaned_text = re.sub(r'(.)\1{2,}', r'\1', text)
    return cleaned_text

def make_columns_unique(columns):
    seen = set()
    unique_columns = []
    for col in columns:
        # 如果列名重复，就添加一个索引
        if col in seen:
            idx = 1
            new_col = f"{col}_{idx}"
            while new_col in seen:
                idx += 1
                new_col = f"{col}_{idx}"
            unique_columns.append(new_col)
            seen.add(new_col)
        else:
            unique_columns.append(col)
            seen.add(col)
    return unique_columns

def process_pdf(file_path):
    full_content = ""
    try:
        # 打开PDF文件
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # 提取当前页面的文字内容
                text = page.extract_text()
                if text:
                    # 清理文本中的重复字符
                    cleaned_text = clean_repeated_characters(text)
                    full_content += cleaned_text + "\n"
                
                # 提取当前页面的表格资料
                table = page.extract_table()
                print(f"table:{table}")
                if table:
                    unique_columns = make_columns_unique(table[0])
                    df = pd.DataFrame(table[1:], columns=unique_columns)  # 创建数据框架，表头为表格的第一行
                    df = df.replace(r'\n', ' ', regex=True)  # 替换数据框架中的换行符为空格
                    table_text = str(df.to_dict(orient='records'))  # 将数据框架转换为字符串
                    full_content += table_text + "\n"
        
        if full_content:
            return full_content.strip()
        else:
            return None

    except Exception as e:
        print(f"An error occurred: {e}")
        return None

def process_and_save_all(file_data, pages_per_split=100):
    temp_folder = "/home/factorytechllm/Laoshifu/PDF_Tentative/"
    # 確保輸出文件夾存在
    if not os.path.exists(temp_folder):
        os.makedirs(temp_folder)
    # 先分割 PDF
    split_files = split_pdf(file_data, temp_folder, pages_per_split)
    
    all_content = ""
    
    try:
        # 處理每個分割後的文件
        for split_file in split_files:
            content = process_pdf(split_file)
            if content:
                all_content += content + "\n"
    
    finally:
        # 確保釋放資源，刪除臨時文件
        for split_file in split_files:
            try:
                os.remove(split_file)
                print(f"已刪除 {split_file}")
            except Exception as e:
                print(f"無法刪除 {split_file}: {e}")
    
    return all_content.strip() if all_content else None

def process_docx(file_data):
    # Convert bytes data to a file-like object
    try:
        file_data_io = BytesIO(file_data)
        doc = docxDocument(file_data_io)
        full_text = [para.text for para in doc.paragraphs]

        for table in doc.tables:
            # column name (first row)
            row_content = {cell.text:[] for cell in table.rows[0].cells}
            # row record
            for row in table.rows[1:]:
                for col, content in zip(row_content.keys(), row.cells):
                    row_content[col].append(content.text)

            for items in zip(*row_content.values()):
                row = dict(zip(row_content.keys(), items))
                full_text.append(''.join(str(row)))
        if full_text:
            return '\n'.join(full_text)
        else:
            return None
    except:
        return None

def process_sheet(file_data, file_name):
    try:
        data = BytesIO(file_data)
        xls = pd.ExcelFile(data)  # 讀取 Excel 檔案
        output_chunks = []

        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name, header=None)

            print(f"Processing sheet: {sheet_name}")

            # 找到第一個包含至少2個非空值的行作為標題行
            title_row_index = None
            first_value_index = None  # 用於儲存第一個非空值的位置
            for index, row in df.iterrows():
                non_empty_values = row.dropna()  # 去除空值
                if len(non_empty_values) >= 2:
                    title_row_index = index
                    first_value_index = row.first_valid_index()  # 獲取第一個非空值的位置
                    break

            if title_row_index is None or first_value_index is None:
                print(f"未找到符合條件的標題行，將提取整個工作表 {sheet_name} 的內容。")
                # 如果找不到標題行，將整個工作表的內容轉為字串
                current_chunk = [f"file_name: {file_name} page: {sheet_name}"]
                for i in range(len(df)):
                    row = df.iloc[i]
                    row_str = []
                    for value in row.dropna():  # 只處理非空值
                        row_str.append(str(value))
                    if row_str:
                        current_chunk.append("\t".join(row_str))

                    # 檢查當前塊的字符數是否超過 700
                    if sum(len(s) for s in current_chunk) > 700:
                        output_chunks.append("\n".join(current_chunk))
                        current_chunk = [f"file_name: {file_name} page: {sheet_name}"]

                if current_chunk:
                    output_chunks.append("\n".join(current_chunk))
                continue  # 跳過當前工作表，繼續處理下一個

            # 獲取標題行的欄位名稱
            titles = df.iloc[title_row_index].dropna().tolist()

            # 計算基準偏移量，用於在資料行中正確對齊
            offset = df.columns.get_loc(first_value_index)

            # 初始化存儲當前工作表的結果字串
            current_chunk = [f"file_name: {file_name} page: {sheet_name}"]

            # 開始從標題行的下一行逐行處理
            for i in range(title_row_index + 1, len(df)):
                row = df.iloc[i]
                if row.isna().all():  # 如果整行都是空的，則終止處理
                    break

                row_str = []
                for j, title in enumerate(titles):
                    value = row.iloc[j + offset]  # 使用偏移量來調整列索引
                    if pd.notna(value):  # 只處理非空值
                        row_str.append(f"{title}: {value}")
                if row_str:
                    row_string = "\t".join(row_str)
                    current_chunk.append(row_string)

                # 檢查當前塊的字符數是否超過 700
                if sum(len(s) for s in current_chunk) > 500:
                    output_chunks.append("\n".join(current_chunk))
                    # 生成新的塊並添加標題信息
                    current_chunk = [f"file_name: {file_name} page: {sheet_name}"]

            # 將最後剩下的塊添加到列表中
            if current_chunk:
                output_chunks.append("\n".join(current_chunk))

        return output_chunks
    except Exception as e:
        return [str(e)]

def create_document(data, collection, index, source, is_sheet=False):
    # 根據是否處理 sheet 資料設定 chunk_size 和 chunk_overlap
    if is_sheet:
        chunk_size = 1000
        chunk_overlap = 0
    else:
        chunk_size = 600
        chunk_overlap = 100

    doc = Document(page_content=data, metadata={"source": source , "updatedAt": datetime.now()})
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap, length_function=len)
    docs = text_splitter.split_documents([doc])

    # Split docs into chunks
    chunks = [docs[i:i + 100] for i in range(0, len(docs), 100)]

    for chunk in chunks:
        # Create the vector store
        vector_store = MongoDBAtlasVectorSearch.from_documents(
            documents=chunk,
            embedding=OpenAIEmbeddings(),
            collection=collection,
            index_name=index
        )
        time.sleep(0.1)  # Sleep for 1 second to avoid overwhelming the system
    
    return "document created successfully"

# if __name__ == "__main__":
#     # 设置本地 PDF 文件路径
#     local_pdf_path = r"C:\Users\pp657\Downloads\先構通識-加班費申請規則.pdf"
    
#     # 读取 PDF 文件的二进制数据
#     with open(local_pdf_path, "rb") as f:
#         pdf_data = f.read()
    
#     # 调用 process_and_save_all 处理 PDF 文件并创建文档
#     processed_content = process_and_save_all(pdf_data, pages_per_split=100)

#     print(processed_content)